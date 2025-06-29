from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.popup import Popup
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.clock import Clock
from kivy.core.window import Window
from kivy.utils import get_color_from_hex
from kivy.graphics import Color, RoundedRectangle, Rectangle
from kivy.uix.gridlayout import GridLayout
import pyttsx3
import speech_recognition as sr
import wikipedia
import requests
import webbrowser
import os
from datetime import datetime
import spacy
from PyDictionary import PyDictionary
import language_tool_python
from tkinter import Tk, filedialog
import docx
import openpyxl
import pandas as pd
from pdfplumber import open as pdf_open
from reportlab.pdfgen import canvas
engine = pyttsx3.init()
Window.size = (400, 600)
OPENWEATHER_API_KEY = "9c2304348e7335364ba0f8e867ab9cb4"
NEWS_API_KEY = "685bbbb435114b4ebe304c3226f508a1"
nlp = spacy.load("en_core_web_sm")
DEFAULT_USERNAME = "Shanmukh"
DEFAULT_PASSWORD = "Jeevan"
current_username = DEFAULT_USERNAME
current_password = DEFAULT_PASSWORD
stop_conversation = False
context = {
    "location": None,
    "topic": None,
    "sentence": None
}
music_path = None
notes = []
search_history = []
dictionary = PyDictionary()
tool = language_tool_python.LanguageTool('en-US')
ACCENT_COLOR = "#FFA726"
BG_COLOR = "#F5F5F5"
TEXT_COLOR = "#333333"
DARK_BG = "#1E1E1E"
DARK_TEXT = "#FFFFFF"
def speak(text):
    print(f"JARVIS: {text}")
    engine.say(text)
    engine.runAndWait()
def listen():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        print("Listening...")
        recognizer.adjust_for_ambient_noise(source)
        audio = recognizer.listen(source, timeout=10)
    try:
        print("Recognizing...")
        query = recognizer.recognize_google(audio, language='en-US')
        print(f"User said: {query}")
        return query.strip().lower()
    except Exception:
        return ""
def get_weather(location):
    url = f"http://api.openweathermap.org/data/2.5/weather?q={location}&appid={OPENWEATHER_API_KEY}&units=metric"
    response = requests.get(url)
    data = response.json()
    if data["cod"] == 200:
        weather = data["weather"][0]["description"]
        temperature = data["main"]["temp"]
        humidity = data["main"]["humidity"]
        wind_speed = data["wind"]["speed"]
        sunrise = datetime.fromtimestamp(data["sys"]["sunrise"]).strftime("%H:%M")
        sunset = datetime.fromtimestamp(data["sys"]["sunset"]).strftime("%H:%M")
        return (f"The weather in {location} is {weather}. "
                f"The temperature is {temperature} degrees Celsius. "
                f"The humidity is {humidity}%. "
                f"The wind speed is {wind_speed} meters per second. "
                f"The sunrise is at {sunrise}, and the sunset is at {sunset}.")
    return f"I couldn't fetch the weather information for {location}."
def search_wikipedia(query):
    try:
        summary = wikipedia.summary(query, sentences=2)
        return f"According to Wikipedia, {summary}"
    except wikipedia.exceptions.DisambiguationError as e:
        options = ", ".join(e.options[:5])
        return f"There are multiple results for '{query}'. Did you mean one of these: {options}?"
    except wikipedia.exceptions.PageError:
        return f"I couldn't find any information about '{query}'. Please try again."
def get_latest_news():
    url = f"https://newsapi.org/v2/top-headlines?country=us&apiKey={NEWS_API_KEY}"
    response = requests.get(url)
    data = response.json()
    if data["status"] == "ok" and data["totalResults"] > 0:
        articles = data["articles"]
        news_headlines = [f"{i + 1}. {article['title']} (Source: {article['source']['name']})" for i, article in
                          enumerate(articles[:5])]
        return "\n".join(news_headlines)
    return "I couldn't fetch the latest news at the moment."
def define_word(word):
    meanings = dictionary.meaning(word)
    if meanings:
        result = f"Definitions for '{word}':\n"
        for part_of_speech, definitions in meanings.items():
            result += f"{part_of_speech.capitalize()}:\n"
            for i, definition in enumerate(definitions, 1):
                result += f"  {i}. {definition}\n"
        return result
    return f"No definitions found for '{word}'."
def explain_grammar(sentence):
    matches = tool.check(sentence)
    if not matches:
        return "The sentence is grammatically correct."
    result = "Grammar suggestions:\n"
    for match in matches:
        result += f"- {match.message}\n"
        result += f"  Suggested correction: {match.replacements}\n"
    return result
def get_medicine_info(medicine_name):
    api_url = f"https://api.fda.gov/drug/label.json?search=openfda.brand_name:{medicine_name}&limit=1"
    try:
        response = requests.get(api_url)
        data = response.json()
        if response.status_code == 200 and "results" in data and len(data["results"]) > 0:
            result = data["results"][0]
            purpose = result.get("purpose", ["Purpose not available"])[0]
            dosage = result.get("dosage_and_administration", ["Dosage not available"])[0]
            warnings = result.get("warnings", ["Warnings not available"])[0]
            return (
                f"Medicine: {medicine_name}\n"
                f"Purpose: {purpose}\n"
                f"Dosage: {dosage}\n"
                f"Warnings: {warnings}"
            )
        else:
            return f"No information found for the medicine '{medicine_name}'. Please check the name and try again."
    except Exception as e:
        return f"An error occurred while fetching medicine data: {str(e)}"
def process_command(command):
    global stop_conversation, context, search_history
    if stop_conversation:
        return "Stopping the conversation. Goodbye!"
    search_history.append(command)
    if "medicine info" in command:
        medicine_name = command.replace("medicine info", "").strip()
        if medicine_name:
            return get_medicine_info(medicine_name)
        else:
            return "Please provide the name of the medicine."
    if "explain grammar" in command:
        context["sentence"] = command.replace("explain grammar", "").strip()
        if context["sentence"]:
            return explain_grammar(context["sentence"])
        else:
            return "Please provide a sentence to explain its grammar."
    elif "city" in command or "country" in command or "continent" in command:
        locations = extract_geographical_entities(command)
        if locations:
            context["location"] = locations[0]
            return f"Location '{context['location']}' stored. How can I assist you with this location?"
    elif "weather" in command and context["location"]:
        return get_weather(context["location"])
    elif "tell me about" in command and context["location"]:
        return search_wikipedia(context["location"])
    elif "search" in command and context["topic"]:
        return search_wikipedia(context["topic"])
    elif "hello" in command or "hi" in command:
        return "Hello! How can I assist you?"
    elif "time" in command:
        return f"The current time is {datetime.now().strftime('%H:%M')}."
    elif "latest news" in command:
        return get_latest_news()
    elif "open google" in command:
        webbrowser.open("https://www.google.com")
        return "Opening Google."
    elif "open youtube" in command:
        webbrowser.open("https://www.youtube.com")
        return "Opening YouTube."
    elif "play music" in command:
        global music_path
        if music_path and os.path.exists(music_path):
            os.startfile(music_path)
            return "Playing music."
        else:
            return "No music file selected. Please add a music file first."
    elif "define" in command:
        word = command.replace("define", "").strip()
        return define_word(word)
    elif "exit" in command or "quit" in command or "stop" in command:
        stop_conversation = True
        return "Goodbye!"
    elif "add note" in command or "save note" in command:
        note = command.replace("add note", "").replace("save note", "").strip()
        if note:
            notes.append(note)
            return f"Note added: {note}"
        else:
            return "Please write something to save as a note."
    elif "view notes" in command:
        if notes:
            return "Your Notes:\n" + "\n".join([f"{i + 1}. {note}" for i, note in enumerate(notes)])
        else:
            return "You have no notes yet."
    else:
        return "I'm sorry, I don't understand that command. Can you clarify?"
def extract_geographical_entities(command):
    doc = nlp(command)
    locations = [ent.text for ent in doc.ents if ent.label_ in ["GPE", "LOC"]]
    return locations
class RoundedButton(Button):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.background_color = get_color_from_hex(ACCENT_COLOR)
        self.color = get_color_from_hex("#FFFFFF")
        self.font_size = "18sp"
        self.size_hint = (1, None)
        self.height = 50
        self.halign = "center"
        self.valign = "middle"
        self.padding = (10, 10)
        with self.canvas.before:
            Color(rgba=get_color_from_hex(ACCENT_COLOR))
            self.rect = RoundedRectangle(size=self.size, pos=self.pos, radius=[20])
        self.bind(pos=self.update_rect, size=self.update_rect)
    def update_rect(self, *args):
        self.rect.pos = self.pos
        self.rect.size = self.size
class VoiceSelectionScreen(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        layout = BoxLayout(orientation='vertical', padding=20, spacing=10)
        title = Label(text="Choose Your Voice", font_size=24, color=get_color_from_hex(TEXT_COLOR))
        layout.add_widget(title)
        voice_options = ["Default", "Female 1", "Female 2", "Male"]
        for idx, name in enumerate(voice_options):
            btn = RoundedButton(text=name)
            btn.bind(on_press=lambda btn, i=idx: self.set_voice(i))
            layout.add_widget(btn)
        back_btn = RoundedButton(text="Back", background_color=get_color_from_hex("#FFA726"))
        back_btn.bind(on_press=lambda x: setattr(self.manager, 'current', 'main'))
        layout.add_widget(back_btn)
        self.add_widget(layout)
    def set_voice(self, voice_index):
        global engine
        voices = engine.getProperty('voices')
        if voice_index < len(voices):
            engine.setProperty('voice', voices[voice_index].id)
        speak("Voice changed successfully.")
        self.manager.current = "main"
class LoginScreen(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.layout = BoxLayout(orientation='vertical', padding=20, spacing=10)
        self.layout.canvas.before.add(Color(rgba=get_color_from_hex(DARK_BG)))
        self.title_label = Label(
            text="LEO - Login",
            font_size="24sp",
            color=get_color_from_hex(DARK_TEXT),
            size_hint_y=None,
            height=50
        )
        self.layout.add_widget(self.title_label)
        self.username_input = TextInput(
            hint_text="Username",
            multiline=False,
            font_size="18sp",
            size_hint_y=None,
            height=50,
            background_color=get_color_from_hex("#424242"),
            foreground_color=get_color_from_hex("#FFFFFF")
        )
        self.layout.add_widget(self.username_input)
        self.password_input = TextInput(
            hint_text="Password",
            password=True,
            multiline=False,
            font_size="18sp",
            size_hint_y=None,
            height=50,
            background_color=get_color_from_hex("#424242"),
            foreground_color=get_color_from_hex("#FFFFFF")
        )
        self.layout.add_widget(self.password_input)
        self.login_button = RoundedButton(text="Login")
        self.login_button.bind(on_press=self.validate_login)
        self.layout.add_widget(self.login_button)
        self.add_widget(self.layout)
    def validate_login(self, instance):
        global current_username, current_password
        username = self.username_input.text.strip()
        password = self.password_input.text.strip()
        if username == current_username and password == current_password:
            self.manager.current = "main"
        else:
            popup = Popup(
                title="Login Failed",
                content=Label(text="Invalid username or password."),
                size_hint=(0.8, 0.4),
                background_color=get_color_from_hex("#FF5252")
            )
            popup.open()
class MainScreen(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.dark_mode = False
        self.layout = BoxLayout(orientation='vertical', padding=10, spacing=10)
        with self.canvas.before:
            Color(rgba=get_color_from_hex(BG_COLOR))
            self.bg_rect = Rectangle(pos=self.pos, size=self.size)
        self.bind(pos=self.update_bg, size=self.update_bg)
        self.title_label = Label(
            text="Welcome to JARVIS!",
            font_size="20sp",
            color=get_color_from_hex(TEXT_COLOR),
            size_hint_y=None,
            height=40
        )
        self.layout.add_widget(self.title_label)
        self.output_label = Label(
            text="Ready to assist you...",
            font_size="16sp",
            color=get_color_from_hex(TEXT_COLOR),
            size_hint_y=None,
            height=100
        )
        self.layout.add_widget(self.output_label)
        self.input_field = TextInput(
            hint_text="Enter your command here...",
            multiline=False,
            font_size="16sp",
            size_hint_y=None,
            height=50,
            background_color=get_color_from_hex("#EEEEEE"),
            foreground_color=get_color_from_hex("#000000")
        )
        self.layout.add_widget(self.input_field)
        self.submit_button = RoundedButton(text="Submit")
        self.submit_button.bind(on_press=self.process_command_wrapper)
        self.layout.add_widget(self.submit_button)
        grid = GridLayout(cols=2, spacing=10, size_hint_y=None)
        grid.bind(minimum_height=grid.setter('height'))
        self.speak_button = RoundedButton(text="Speak Mode")
        self.speak_button.bind(on_press=self.start_speak_mode)
        grid.add_widget(self.speak_button)
        self.note_add_button = RoundedButton(text="Add Note")
        self.note_add_button.bind(on_press=self.add_note)
        grid.add_widget(self.note_add_button)
        self.view_notes_button = RoundedButton(text="View Notes")
        self.view_notes_button.bind(on_press=self.view_notes)
        grid.add_widget(self.view_notes_button)
        self.play_music_button = RoundedButton(text="Play Music")
        self.play_music_button.bind(on_press=self.play_music)
        grid.add_widget(self.play_music_button)
        self.voice_button = RoundedButton(text="Change Voice")
        self.voice_button.bind(on_press=lambda x: setattr(self.manager, 'current', 'voice_selection'))
        grid.add_widget(self.voice_button)
        self.theme_button = RoundedButton(text="Toggle Theme")
        self.theme_button.bind(on_press=self.toggle_theme)
        grid.add_widget(self.theme_button)
        self.add_music_button = RoundedButton(text="Add Music")
        self.add_music_button.bind(on_press=self.add_music_file)
        grid.add_widget(self.add_music_button)
        self.stop_button = RoundedButton(text="Stop", background_color=get_color_from_hex("#FF5252"))
        self.stop_button.bind(on_press=self.stop_conversation_func)
        grid.add_widget(self.stop_button)
        self.dashboard_button = RoundedButton(text="Dashboard")
        self.dashboard_button.bind(on_press=lambda x: setattr(self.manager, 'current', 'dashboard'))
        grid.add_widget(self.dashboard_button)
        self.file_converter_button = RoundedButton(text="File Converter")
        self.file_converter_button.bind(on_press=lambda x: setattr(self.manager, 'current', 'file_converter'))
        grid.add_widget(self.file_converter_button)
        self.layout.add_widget(grid)
        self.add_widget(self.layout)
    def update_bg(self, *args):
        self.bg_rect.pos = self.pos
        self.bg_rect.size = self.size
    def toggle_theme(self, instance):
        self.dark_mode = not self.dark_mode
        if self.dark_mode:
            self.layout.canvas.before.clear()
            with self.layout.canvas.before:
                Color(rgba=get_color_from_hex(DARK_BG))
                self.bg_rect = Rectangle(pos=self.pos, size=self.size)
            self.output_label.color = get_color_from_hex(DARK_TEXT)
            self.title_label.color = get_color_from_hex(DARK_TEXT)
            self.theme_button.text = "Light Mode"
        else:
            self.layout.canvas.before.clear()
            with self.layout.canvas.before:
                Color(rgba=get_color_from_hex(BG_COLOR))
                self.bg_rect = Rectangle(pos=self.pos, size=self.size)
            self.output_label.color = get_color_from_hex(TEXT_COLOR)
            self.title_label.color = get_color_from_hex(TEXT_COLOR)
            self.theme_button.text = "Dark Mode"

    def process_command_wrapper(self, instance):
        command = self.input_field.text.strip().lower()
        if command:
            response = process_command(command)
            self.output_label.text = f"JARVIS: {response}"
            speak(response)
            self.input_field.text = ""

    def start_speak_mode(self, instance):
        def listen_and_process(dt):
            global stop_conversation
            if stop_conversation:
                return False
            command = listen()
            if command:
                response = process_command(command)
                self.output_label.text = f"JARVIS: {response}"
                speak(response)
        Clock.schedule_interval(listen_and_process, 1)
    def stop_conversation_func(self, instance):
        global stop_conversation
        stop_conversation = True
        self.output_label.text = "Stopping the conversation. Goodbye!"
        speak("Stopping the conversation. Goodbye!")
    def add_music_file(self, instance):
        global music_path
        root = Tk()
        root.withdraw()
        file_path = filedialog.askopenfilename(
            title="Select a Music File",
            filetypes=[("Audio Files", "*.mp3 *.wav"), ("All Files", "*.*")]
        )
        if file_path:
            music_path = file_path
            self.output_label.text = f"Music file added: {os.path.basename(music_path)}"
        else:
            self.output_label.text = "No music file selected."

    def play_music(self, instance):
        global music_path
        if music_path and os.path.exists(music_path):
            os.startfile(music_path)
            self.output_label.text = "Playing music..."
        else:
            self.output_label.text = "No music file selected. Please add a music file first."
    def add_note(self, instance):
        note = self.input_field.text.strip()
        if note:
            notes.append(note)
            self.output_label.text = f"Note added: {note}"
            self.input_field.text = ""
        else:
            self.output_label.text = "Please enter a note before saving."
    def view_notes(self, instance):
        if notes:
            notes_list = "\n".join([f"{i + 1}. {note}" for i, note in enumerate(notes)])
            self.output_label.text = f"Your Notes:\n{notes_list}"
        else:
            self.output_label.text = "You have no notes yet."
class DashboardScreen(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.layout = BoxLayout(orientation='vertical', padding=20, spacing=10)
        self.title = Label(text="Dashboard", font_size='24sp', color=get_color_from_hex(TEXT_COLOR))
        self.layout.add_widget(self.title)
        self.user_label = Label(text=f"Logged in as: {current_username}", font_size='18sp',
                                color=get_color_from_hex(TEXT_COLOR))
        self.layout.add_widget(self.user_label)
        self.history_label = Label(text="Recent Searches:", font_size='18sp', color=get_color_from_hex(TEXT_COLOR))
        self.layout.add_widget(self.history_label)
        self.history_output = Label(text=self.format_searches(), font_size='14sp', color=get_color_from_hex(TEXT_COLOR))
        self.layout.add_widget(self.history_output)
        self.reset_button = RoundedButton(text="Reset Password")
        self.reset_button.bind(on_press=self.reset_password)
        self.layout.add_widget(self.reset_button)
        self.logout_button = RoundedButton(text="Logout")
        self.logout_button.bind(on_press=self.logout)
        self.layout.add_widget(self.logout_button)
        self.back_button = RoundedButton(text="Back")
        self.back_button.bind(on_press=lambda x: setattr(self.manager, 'current', 'main'))
        self.layout.add_widget(self.back_button)
        self.add_widget(self.layout)
    def format_searches(self):
        if not search_history:
            return "No recent searches."
        return "\n".join([f"{i + 1}. {cmd}" for i, cmd in enumerate(search_history[-5:])])
    def reset_password(self, instance):
        box = BoxLayout(orientation='vertical')
        pwd_input = TextInput(hint_text="Enter new password", password=True)
        confirm_btn = Button(text="Confirm")
        box.add_widget(pwd_input)
        box.add_widget(confirm_btn)
        popup = Popup(title="Reset Password", content=box, size_hint=(0.8, 0.4))
        def update_password(instance):
            new_pass = pwd_input.text.strip()
            if new_pass:
                global current_password
                current_password = new_pass
                popup.dismiss()
                success_popup = Popup(title="Success", content=Label(text="Password updated successfully."),
                                      size_hint=(0.6, 0.3))
                success_popup.open()
        confirm_btn.bind(on_press=update_password)
        popup.open()
    def logout(self, instance):
        self.manager.current = "login"
class FileConverterScreen(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        layout = BoxLayout(orientation='vertical', padding=20, spacing=10)
        title = Label(
            text="File Converter",
            font_size='24sp',
            color=get_color_from_hex(TEXT_COLOR),
            size_hint_y=None,
            height=40
        )
        layout.add_widget(title)
        self.output_label = Label(
            text="Select a file and choose conversion type.",
            font_size='16sp',
            color=get_color_from_hex(TEXT_COLOR),
            size_hint_y=None,
            height=80
        )
        layout.add_widget(self.output_label)
        btn_layout = BoxLayout(orientation='vertical', spacing=10)
        buttons = [
            ("Convert PDF to Word", self.pdf_to_word),
            ("Convert Word to PDF", self.word_to_pdf),
            ("Convert Excel to Word", self.excel_to_word),
            ("Convert Excel to PDF", self.excel_to_pdf),
            ("Convert Word to Excel", self.word_to_excel),
        ]
        for text, func in buttons:
            btn = RoundedButton(text=text)
            btn.bind(on_press=func)
            btn_layout.add_widget(btn)
        back_btn = RoundedButton(text="Back")
        back_btn.bind(on_press=lambda x: setattr(self.manager, 'current', 'main'))
        btn_layout.add_widget(back_btn)
        layout.add_widget(btn_layout)
        self.add_widget(layout)
    def show_output(self, message):
        self.output_label.text = message
    def choose_file(self, filetypes=[("All Files", "*.*")]):
        Tk().withdraw()
        path = filedialog.askopenfilename(filetypes=filetypes)
        return path
    def save_file_dialog(self, defaultextension=".docx"):
        Tk().withdraw()
        path = filedialog.asksaveasfilename(defaultextension=defaultextension)
        return path
    def open_with_default_app(self, filepath):
        try:
            if os.name == 'nt':  # Windows
                os.startfile(filepath)
            elif os.uname().sysname == 'Darwin':  # macOS
                subprocess.call(['open', filepath])
            else:  # Linux
                subprocess.call(['xdg-open', filepath])
            self.show_output(f"Converted file opened:\n{filepath}")
        except Exception as e:
            self.show_output(f"Could not open file: {e}")
    def pdf_to_word(self, instance=None):
        path = self.choose_file([("PDF Files", "*.pdf")])
        if not path:
            self.show_output("No file selected.")
            return
        output_path = self.save_file_dialog(".docx")
        try:
            doc = docx.Document()
            with pdf_open(path) as pdf_file:
                for page in pdf_file.pages:
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)
            doc.save(output_path)
            self.open_with_default_app(output_path)
        except Exception as e:
            self.show_output(f"Error: {str(e)}")
    def word_to_pdf(self, instance=None):
        path = self.choose_file([("Word Files", "*.docx")])
        if not path:
            self.show_output("No file selected.")
            return
        output_path = self.save_file_dialog(".pdf")
        try:
            c = canvas.Canvas(output_path)
            doc = docx.Document(path)
            y = 750
            for para in doc.paragraphs:
                c.drawString(50, y, para.text)
                y -= 15
                if y < 50:
                    c.showPage()
                    y = 750
            c.save()
            self.open_with_default_app(output_path)
        except Exception as e:
            self.show_output(f"Error: {str(e)}")
    def excel_to_word(self, instance=None):
        path = self.choose_file([("Excel Files", "*.xlsx")])
        if not path:
            self.show_output("No file selected.")
            return
        output_path = self.save_file_dialog(".docx")
        try:
            wb = openpyxl.load_workbook(path)
            ws = wb.active
            doc = docx.Document()
            table = doc.add_table(rows=ws.max_row, cols=ws.max_column)

            for r_idx, row in enumerate(ws.iter_rows()):
                for c_idx, cell in enumerate(row):
                    table.cell(r_idx, c_idx).text = str(cell.value)
            doc.save(output_path)
            self.open_with_default_app(output_path)
        except Exception as e:
            self.show_output(f"Error: {str(e)}")
    def excel_to_pdf(self, instance=None):
        path = self.choose_file([("Excel Files", "*.xlsx")])
        if not path:
            self.show_output("No file selected.")
            return
        output_path = self.save_file_dialog(".pdf")
        try:
            df = pd.read_excel(path)
            c = canvas.Canvas(output_path)
            y = 750
            for _, row in df.iterrows():
                line = " | ".join(str(x) for x in row.values)
                c.drawString(50, y, line)
                y -= 15
                if y < 50:
                    c.showPage()
                    y = 750
            c.save()
            self.open_with_default_app(output_path)
        except Exception as e:
            self.show_output(f"Error: {str(e)}")
    def word_to_excel(self, instance=None):
        path = self.choose_file([("Word Files", "*.docx")])
        if not path:
            self.show_output("No file selected.")
            return
        output_path = self.save_file_dialog(".xlsx")
        try:
            doc = docx.Document(path)
            data = []
            for para in doc.paragraphs:
                if para.text.strip():
                    data.append([para.text])
            df = pd.DataFrame(data, columns=["Text"])
            df.to_excel(output_path, index=False)
            self.open_with_default_app(output_path)
        except Exception as e:
            self.show_output(f"Error: {str(e)}")
class JarvisApp(App):
    def build(self):
        sm = ScreenManager()
        sm.add_widget(LoginScreen(name="login"))
        sm.add_widget(MainScreen(name="main"))
        sm.add_widget(VoiceSelectionScreen(name="voice_selection"))
        sm.add_widget(DashboardScreen(name="dashboard"))
        sm.add_widget(FileConverterScreen(name="file_converter"))  # New screen added
        return sm
if __name__ == "__main__":
    JarvisApp().run()
